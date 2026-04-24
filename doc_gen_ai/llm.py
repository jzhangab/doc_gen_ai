import io
import json
import logging
import time
from datetime import datetime

from . import config

logger = logging.getLogger(__name__)

_MAX_RETRIES = 3
_RETRY_DELAY = 5  # seconds between retries


def _trunc(text: str) -> str:
    return text[:config._MAX_CHARS] + "\n...[truncated]" if len(text) > config._MAX_CHARS else text


def _llm_call(messages: list, connection_id: str = None) -> str:
    import dataiku
    conn = connection_id or config.DEFAULT_LLM_CONNECTION_ID
    api_client = dataiku.api_client()
    project = api_client.get_project(dataiku.default_project_key())
    llm = project.get_llm(conn)

    last_exc = None
    for attempt in range(1, _MAX_RETRIES + 1):
        try:
            completion = llm.new_completion()
            for msg in messages:
                completion.with_message(msg["content"], msg["role"])
            try:
                completion.with_max_output_tokens(8192)
            except Exception:
                pass
            return completion.execute().text
        except Exception as exc:
            last_exc = exc
            if attempt < _MAX_RETRIES:
                logger.warning(
                    "LLM call failed (attempt %d/%d): %s — retrying in %ds…",
                    attempt, _MAX_RETRIES, exc, _RETRY_DELAY,
                )
                print(f"      [retry {attempt}/{_MAX_RETRIES}] LLM error: {exc} — retrying in {_RETRY_DELAY}s…")
                time.sleep(_RETRY_DELAY)
            else:
                logger.error("LLM call failed after %d attempts: %s", _MAX_RETRIES, exc)
    raise last_exc


def _llm_json(messages: list, connection_id: str = None) -> dict:
    last_exc = None
    for attempt in range(1, _MAX_RETRIES + 1):
        raw = _llm_call(messages, connection_id=connection_id).strip()
        if raw.startswith("```"):
            raw = "\n".join(
                ln for ln in raw.splitlines() if not ln.strip().startswith("```")
            ).strip()
        try:
            return json.loads(raw)
        except json.JSONDecodeError as exc:
            last_exc = exc
            if attempt < _MAX_RETRIES:
                logger.warning(
                    "JSON decode failed (attempt %d/%d): %s — retrying…",
                    attempt, _MAX_RETRIES, exc,
                )
                print(f"      [retry {attempt}/{_MAX_RETRIES}] JSON parse error: {exc} — retrying…")
            else:
                logger.error("JSON decode failed after %d attempts: %s", _MAX_RETRIES, exc)
    raise last_exc


def select_relevant_templates(filenames: list, doc_type: str, connection_id: str = None) -> str:
    """Return the single filename most consistent with doc_type."""
    file_list = "\n".join(f"- {f}" for f in filenames)
    result = _llm_json([
        {
            "role": "system",
            "content": (
                "You are an expert in ISO 13485 software V&V documentation. "
                "Return only valid JSON."
            ),
        },
        {
            "role": "user",
            "content": (
                f'I need to generate a "{doc_type}" document.\n\n'
                f"Available files in the templates folder:\n{file_list}\n\n"
                f'Choose the single file whose name most closely matches a "{doc_type}". '
                "Consider abbreviations, version numbers, and partial name matches. "
                'Return JSON: {"selected": "filename"}'
            ),
        },
    ], connection_id=connection_id)
    return result.get("selected", "")


def extract_writing_context(example_texts: list, connection_id: str = None) -> str:
    """Synthesise a writing style guide from context example documents.

    Extracts tone, language patterns, and phrasing conventions without
    reproducing any specific content from the examples.
    Returns a plain-text style guide (300-500 words).
    """
    combined = "\n\n---\n\n".join(
        f"[Example {i+1}]\n{_trunc(t)}" for i, t in enumerate(example_texts)
    )
    return _llm_call([
        {
            "role": "system",
            "content": (
                "You are an expert in technical writing for ISO 13485 regulated software. "
                "Analyse example documents and produce a writing style guide. "
                "Focus entirely on style, tone, and language patterns. "
                "Do not reproduce, summarise, or reference any specific content from the examples."
            ),
        },
        {
            "role": "user",
            "content": (
                "Analyse these example technical documents and produce a concise writing style guide "
                "for use when authoring new documents of the same type.\n\n"
                f"{combined}\n\n"
                "Capture:\n"
                "- Overall tone and register (formality level, active vs passive voice, person)\n"
                "- How technical claims, rationale, and evidence are typically expressed\n"
                "- Characteristic sentence structures, transitions, and paragraph patterns\n"
                "- How requirements, standards, and compliance obligations are referenced\n"
                "- Domain-specific terminology and phrasing conventions\n"
                "- What to avoid (verbosity, hedging language, generic statements)\n\n"
                "Write this as a practical style guide for a technical writer. "
                "Do NOT quote, paraphrase, or reference any specific facts, names, systems, "
                "or details from the example documents — only the writing patterns matter. "
                "Keep it to 400-500 words."
            ),
        },
    ], connection_id=connection_id)


def discover_template_structure(template_texts: list, doc_type: str, connection_id: str = None) -> dict:
    """Analyse example documents and extract section structure, style, and regulatory language."""
    combined = "\n\n---\n\n".join(
        f"[Example {i+1}]\n{_trunc(t)}" for i, t in enumerate(template_texts)
    )
    return _llm_json([
        {
            "role": "system",
            "content": (
                "You are an expert in ISO 13485 software V&V documentation. "
                "Return only valid JSON."
            ),
        },
        {
            "role": "user",
            "content": (
                f'Analyse these example "{doc_type}" documents and extract their structure.\n\n'
                f"{combined}\n\n"
                "Rules for the sections list:\n"
                "- Each section heading must be unique — no two entries may cover the same topic.\n"
                "- If the template has overlapping or near-duplicate sections, merge them into one.\n"
                "- Do not include a Table of Contents section.\n\n"
                "Return JSON:\n"
                '{"sections":[{"heading":"...","description":"...","required_elements":["..."]}],'
                '"style_notes":"...","regulatory_language":["..."]}'
            ),
        },
    ], connection_id=connection_id)


def deep_research(doc_texts: list, doc_type: str, connection_id: str = None) -> dict:
    """Extract comprehensive project intelligence needed to write doc_type."""
    combined = "\n\n---\n\n".join(
        f"[Document {i+1}]\n{_trunc(t)}" for i, t in enumerate(doc_texts)
    )
    return _llm_json([
        {
            "role": "system",
            "content": (
                "You are an expert technical analyst for ISO 13485 regulated software. "
                "Conduct deep research on the provided project documentation and extract "
                "every piece of information that will be needed to author a compliance document. "
                "Return only valid JSON."
            ),
        },
        {
            "role": "user",
            "content": (
                f'Conduct deep research on these project documents to extract everything '
                f'needed to write a "{doc_type}" per ISO 13485.\n\n'
                f"{combined}\n\n"
                "Return a comprehensive JSON object:\n"
                '{\n'
                '  "system_name": "...",\n'
                '  "system_description": "...",\n'
                '  "purpose": "...",\n'
                '  "requirements": [],\n'
                '  "design_specifications": {},\n'
                '  "testing_methodology": {},\n'
                '  "infrastructure": {},\n'
                '  "stakeholders": [],\n'
                '  "regulatory_context": "...",\n'
                '  "key_decisions": [],\n'
                '  "assumptions": [],\n'
                '  "constraints": [],\n'
                '  "discovered_artifacts": {}\n'
                '}'
            ),
        },
    ], connection_id=connection_id)


def generate_section(
    doc_type: str, section: dict, research: dict,
    style: str, reg_lang: list,
    writing_context: str = None,
    connection_id: str = None,
) -> str:
    required = "\n".join(f"- {e}" for e in section.get("required_elements", []))
    reg = "\n".join(f"- {p}" for p in reg_lang[:10])
    writing_context_block = (
        f"\nWriting style reference — use this to guide tone and language only. "
        f"These are patterns from similar documents; do not copy their content literally:\n"
        f"{writing_context}\n"
        if writing_context else ""
    )
    return _llm_call([
        {
            "role": "system",
            "content": (
                "You are an expert technical writer for ISO 13485 regulated "
                "software V&V documentation. "
                "Write concisely: every sentence must earn its place. "
                "State facts directly — no preamble, no filler, no restating what the heading already says. "
                "Aim for the minimum word count that fully satisfies the section's purpose.\n\n"
                "Formatting rules for Microsoft Word rendering:\n"
                "- PARAGRAPHS are the default. Use full sentences; do not pad them.\n"
                "- Use bullet lists ONLY when items are genuinely enumerable with no natural prose "
                "flow — e.g. a set of 4+ parallel, discrete items. Never convert a sentence into "
                "a bullet just to break it up.\n"
                "- Use numbered lists ONLY for sequential steps or ranked items.\n"
                "- Use ## for sub-headings, ### for sub-sub-headings\n"
                "- Use **text** for bold\n"
                "- Use pipe syntax for tables: | Col1 | Col2 |\\n|---|---|\\n| val | val |\n"
                "- Separate paragraphs and blocks with a blank line\n"
                "Do not use any other markdown syntax."
            ),
        },
        {
            "role": "user",
            "content": (
                f'Write the "{section["heading"]}" section for a "{doc_type}" '
                f"document per ISO 13485.\n\n"
                f"Purpose: {section['description']}\n"
                f"Required elements:\n{required}\n\n"
                f"Project research:\n{json.dumps(research, indent=2)[:6000]}\n\n"
                f"Template style notes: {style}\n"
                f"Regulatory language to use:\n{reg}\n"
                f"{writing_context_block}\n"
                "Write professional, complete content. Do not repeat the section heading."
            ),
        },
    ], connection_id=connection_id)


_GDP_RULES = """\
Good Documentation Practice (GDP) rules to enforce:
- ATTRIBUTABLE: All claims, decisions, and results must be attributable to a defined role, \
system, or process — not left anonymous or vague.
- ACCURATE & UNAMBIGUOUS: No speculative language ("may", "might", "could", "should consider"). \
Statements must be definitive and precise.
- COMPLETE: No placeholders, TBDs, bracketed gaps, or omitted required elements. Every required \
element of the section must be substantively addressed.
- CONSISTENT: Terminology, abbreviations, and system names must be used consistently throughout. \
Acronyms must be defined on first use.
- TRACEABLE: References to requirements, standards (e.g. ISO 13485, IEC 62304), or other \
documents must be explicit and specific, not generic.
- LEGIBLE & CLEAR: No ambiguous pronouns, run-on logic, or contradictory statements.
- CONTEMPORANEOUS CONTEXT: Dates and version references must be present where relevant.
"""


def gdp_check(doc_texts: list, connection_id: str = None) -> list:
    """Run a GDP-only audit on one or more document texts.

    Returns a list of issues: [{file, rule, location, description}]
    """
    combined = "\n\n---\n\n".join(
        f"[Document {i+1}]\n{_trunc(t)}" for i, t in enumerate(doc_texts)
    )
    result = _llm_json([
        {
            "role": "system",
            "content": (
                "You are a GDP auditor for ISO 13485 regulated software documentation. "
                "Identify every Good Documentation Practice violation in the provided documents. "
                "Return only valid JSON."
            ),
        },
        {
            "role": "user",
            "content": (
                "Audit these documents for GDP violations.\n\n"
                f"{combined}\n\n"
                f"GDP rules to check against:\n{_GDP_RULES}\n"
                "For each violation return a JSON object. "
                "Keep every string field under 20 words. "
                "Group multiple violations of the same rule in the same section into one entry.\n"
                '{"issues": [{"document": "Document N", "rule": "ATTRIBUTABLE|ACCURATE|COMPLETE|'
                'CONSISTENT|TRACEABLE|LEGIBLE|CONTEMPORANEOUS", '
                '"location": "section or paragraph reference", "description": "concise description"}]}\n'
                "Return an empty list if no violations found."
            ),
        },
    ], connection_id=connection_id)
    return result.get("issues", [])


def critique_document(doc_type: str, sections: list, connection_id: str = None) -> list:
    """Scan all sections and return a list of issues to fix.

    Each issue: {index, heading, type, description}
    type: "duplicate" | "formatting" | "incomplete" | "gdp"
    """
    overview = "\n\n".join(
        f"[Section {i+1}: {h}]\n{c[:600]}{'…' if len(c) > 600 else ''}"
        for i, (h, c) in enumerate(sections)
    )
    result = _llm_json([
        {
            "role": "system",
            "content": (
                "You are a strict document quality reviewer for ISO 13485 V&V documentation "
                "with expertise in Good Documentation Practice (GDP). "
                "Identify every quality issue in the document sections. Return only valid JSON."
            ),
        },
        {
            "role": "user",
            "content": (
                f'Review these sections of a "{doc_type}" document.\n\n'
                f"{overview}\n\n"
                "Check for ALL of the following issue types:\n\n"
                "1. DUPLICATE — content that substantially repeats another section\n"
                "2. FORMATTING — broken or inconsistent formatting "
                "(raw markdown symbols, mismatched list styles)\n"
                "3. INCOMPLETE — generic, placeholder-only, or missing required elements\n"
                f"4. GDP — any violation of the following rules:\n{_GDP_RULES}\n"
                "Return JSON — empty list if no issues. "
                "Report every GDP violation as a separate issue entry:\n"
                '{"issues": [{"index": 0, "heading": "...", '
                '"type": "duplicate|formatting|incomplete|gdp", "description": "..."}]}'
            ),
        },
    ], connection_id=connection_id)
    return result.get("issues", [])


def deduplicate_sections(doc_type: str, sections: list, connection_id: str = None) -> list:
    """Final deduplication pass: returns sections with redundant entries removed.

    Asks the LLM which section indices to drop, then filters the list.
    """
    if len(sections) <= 1:
        return sections

    overview = "\n\n".join(
        f"[{i}] {h}\n{c[:1500]}{'…' if len(c) > 1500 else ''}"
        for i, (h, c) in enumerate(sections)
    )
    result = _llm_json([
        {
            "role": "system",
            "content": (
                "You are a document editor whose only job is to eliminate redundancy. "
                "Be aggressive: when in doubt, remove. Return only valid JSON."
            ),
        },
        {
            "role": "user",
            "content": (
                f'These are the {len(sections)} sections of a "{doc_type}" document:\n\n'
                f"{overview}\n\n"
                "Mark a section for removal if ANY of the following are true:\n"
                "1. Its heading is semantically equivalent to another section's heading "
                "(e.g. 'Scope' and 'Document Scope', 'Overview' and 'Introduction').\n"
                "2. Its content substantially overlaps with or is contained within another section.\n"
                "3. It repeats information already present elsewhere in the document.\n\n"
                "When two sections overlap, keep the one that is more complete or better positioned; "
                "mark the other for removal. When in doubt, remove.\n\n"
                "Return ONLY the integer indices to remove. Empty list if nothing to remove:\n"
                '{"remove": [2, 5]}'
            ),
        },
    ], connection_id=connection_id)

    to_remove = set(result.get("remove", []))
    return [(h, c) for i, (h, c) in enumerate(sections) if i not in to_remove]


def fix_section_content(
    doc_type: str, heading: str, content: str, issue: str,
    other_headings: list = None, connection_id: str = None,
) -> str:
    """Rewrite a section to address a specific quality issue.

    other_headings lists the headings of every other section in the document
    so the fixer does not introduce content that duplicates those sections.
    """
    scope_note = ""
    if other_headings:
        scope_note = (
            "\n\nIMPORTANT — the following sections already exist in this document. "
            "Do NOT duplicate their content here; keep this section strictly within its own scope:\n"
            + "\n".join(f"- {h}" for h in other_headings)
        )
    return _llm_call([
        {
            "role": "system",
            "content": (
                "You are an expert technical writer for ISO 13485 V&V documentation "
                "with expertise in Good Documentation Practice (GDP). "
                "Fix the provided section content according to the issue described. "
                "Ensure the corrected content satisfies all GDP rules:\n"
                f"{_GDP_RULES}\n"
                "Write concisely: state facts directly, no preamble or filler. "
                "Every sentence must serve a purpose. "
                "Use bullet lists only for genuinely enumerable parallel items (4+), "
                "numbered lists only for sequential steps. "
                "Format using: ## sub-headings, **bold**, pipe tables. "
                "Return only the corrected section content."
            ),
        },
        {
            "role": "user",
            "content": (
                f'Fix the "{heading}" section of a "{doc_type}" document.\n\n'
                f"Issue: {issue}\n\n"
                f"Current content:\n{content}"
                f"{scope_note}\n\n"
                "Return the corrected content only. Do not repeat the section heading."
            ),
        },
    ], connection_id=connection_id)


def generate_summary(doc_type: str, sections: list, connection_id: str = None) -> str:
    """Generate a 1-page executive summary with key features and insights."""
    overview = "\n\n".join(
        f"## {heading}\n{content[:1000]}{'…' if len(content) > 1000 else ''}"
        for heading, content in sections
    )
    return _llm_call([
        {
            "role": "system",
            "content": (
                "You are an expert technical writer for ISO 13485 regulated software V&V documentation. "
                "Write concisely and precisely. Every sentence must serve a purpose. "
                "No filler, no repetition, no preamble.\n\n"
                "Formatting rules for Microsoft Word rendering:\n"
                "- Use bullet lists for genuinely enumerable parallel items.\n"
                "- Use **text** for bold.\n"
                "- Use ## for sub-headings.\n"
                "- Separate blocks with a blank line.\n"
                "Do not use any other markdown syntax."
            ),
        },
        {
            "role": "user",
            "content": (
                f'Write a 1-page executive summary for the following "{doc_type}" document.\n\n'
                f"Document sections:\n{overview}\n\n"
                "Requirements:\n"
                "- 350–450 words maximum (one page)\n"
                "- Open with 2–3 sentences identifying the system and its purpose\n"
                "- List key features, capabilities, or specifications covered\n"
                "- Highlight critical compliance decisions, risks, or constraints\n"
                "- Note any significant assumptions or dependencies\n\n"
                "Synthesise insights — do not reproduce full section content. "
                "Do not include a title or top-level heading."
            ),
        },
    ], connection_id=connection_id)


def assemble_summary_docx(doc_type: str, summary_text: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    title = doc.add_heading(f"{doc_type} — Executive Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(
        f"Standard: ISO 13485  |  Generated: {datetime.now().strftime('%Y-%m-%d')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        run.font.size = Pt(10)

    doc.add_paragraph()
    _render_content(doc, summary_text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_TOC_SECTION_NAMES = {"table of contents", "contents", "toc"}


def assemble_docx(doc_type: str, sections: list) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # ── Cover page ────────────────────────────────────────────────────────────
    title = doc.add_heading(doc_type, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(
        f"Standard: ISO 13485  |  Generated: {datetime.now().strftime('%Y-%m-%d')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        run.font.size = Pt(10)

    doc.add_page_break()

    # ── Table of Contents ─────────────────────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)

    visible_sections = [
        (i, h, c) for i, (h, c) in enumerate(sections, 1)
        if h.strip().lower() not in _TOC_SECTION_NAMES
    ]

    toc_table = doc.add_table(rows=1 + len(visible_sections), cols=2)
    toc_table.style = "Table Grid"

    # Column widths: narrow number col, wide heading col
    for row in toc_table.rows:
        row.cells[0].width = Inches(0.55)
        row.cells[1].width = Inches(5.2)

    # Header row
    hdr_cells = toc_table.rows[0].cells
    for cell, text in zip(hdr_cells, ("#", "Section")):
        cell.text = text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Dark blue fill
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E3A5F")
        tc_pr.append(shd)

    # Data rows
    for row_idx, (num, heading, _) in enumerate(visible_sections, 1):
        cells = toc_table.rows[row_idx].cells
        cells[0].text = str(num)
        cells[1].text = heading
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Alternating light blue tint on even rows
        if row_idx % 2 == 0:
            for cell in cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EBF2FA")
                tc_pr.append(shd)

    doc.add_paragraph()  # spacing after table
    doc.add_page_break()

    # ── Sections (skip any TOC section returned by the LLM) ──────────────────
    for num, heading, content in visible_sections:
        doc.add_heading(f"{num}. {heading}", level=1)
        _render_content(doc, content)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_content(doc, content: str) -> None:
    import re
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_runs(para, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            para = doc.add_paragraph(style="List Number")
            _add_runs(para, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # Markdown table — collect all consecutive pipe rows
        if stripped.startswith("|"):
            i = _render_table(doc, lines, i)
            continue

        # Plain paragraph — accumulate consecutive non-special lines
        block = []
        while i < len(lines):
            s = lines[i].strip()
            if not s:
                break
            if s.startswith(("## ", "### ", "- ", "* ", "|")) or re.match(r"^\d+\.\s", s):
                break
            block.append(s)
            i += 1
        if block:
            para = doc.add_paragraph()
            _add_runs(para, " ".join(block))


def _render_table(doc, lines: list, start: int) -> int:
    import re
    from docx.shared import RGBColor

    # Collect consecutive pipe-delimited rows
    raw_rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw_rows.append(lines[i].strip())
        i += 1

    # Parse cells, skipping separator rows (|---|---|)
    parsed = []
    for row in raw_rows:
        if re.match(r"^\|[-| :]+\|$", row):
            continue
        cells = [c.strip() for c in row.strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return i

    num_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row_cells in enumerate(parsed):
        for c_idx in range(num_cols):
            text = row_cells[c_idx] if c_idx < len(row_cells) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            _add_runs(para, text)
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True

    # Blank paragraph after table for spacing
    doc.add_paragraph()
    return i


def _add_runs(para, text: str) -> None:
    import re
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            para.add_run(part[2:-2]).bold = True
        else:
            para.add_run(part)
