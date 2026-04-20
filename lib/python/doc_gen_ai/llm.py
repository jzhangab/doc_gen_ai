import io
import json
import logging
from datetime import datetime

from . import config

logger = logging.getLogger(__name__)


def _trunc(text: str) -> str:
    return text[:config._MAX_CHARS] + "\n...[truncated]" if len(text) > config._MAX_CHARS else text


def _llm_call(messages: list, connection_id: str = None) -> str:
    import dataiku
    conn = connection_id or config.DEFAULT_LLM_CONNECTION_ID
    api_client = dataiku.api_client()
    project = api_client.get_project(dataiku.default_project_key())
    llm = project.get_llm(conn)
    completion = llm.new_completion()
    for msg in messages:
        completion.with_message(msg["content"], msg["role"])
    try:
        completion.with_max_output_tokens(8192)
    except Exception:
        pass
    return completion.execute().text


def _llm_json(messages: list, connection_id: str = None) -> dict:
    raw = _llm_call(messages, connection_id=connection_id).strip()
    if raw.startswith("```"):
        raw = "\n".join(
            ln for ln in raw.splitlines() if not ln.strip().startswith("```")
        ).strip()
    return json.loads(raw)


def _discover_template(template_texts: list, doc_type: str, connection_id: str = None) -> dict:
    combined = "\n\n---\n\n".join(
        f"[Example {i+1}]\n{_trunc(t)}" for i, t in enumerate(template_texts)
    )
    return _llm_json([
        {
            "role": "system",
            "content": (
                "You are an expert in ISO 13485 software V&V documentation. "
                "Return only valid JSON."
            ),
        },
        {
            "role": "user",
            "content": (
                f'Analyze these example "{doc_type}" documents and extract their template structure.\n\n'
                f"{combined}\n\n"
                "Return JSON:\n"
                '{"sections":[{"heading":"...","description":"...","required_elements":["..."]}],'
                '"style_notes":"...","regulatory_language":["..."]}'
            ),
        },
    ], connection_id=connection_id)


def _research_inputs(input_texts: list, doc_type: str, connection_id: str = None) -> dict:
    combined = "\n\n---\n\n".join(
        f"[Input {i+1}]\n{_trunc(t)}" for i, t in enumerate(input_texts)
    )
    return _llm_json([
        {
            "role": "system",
            "content": (
                "You are an expert technical analyst for ISO 13485 regulated software. "
                "Return only valid JSON."
            ),
        },
        {
            "role": "user",
            "content": (
                f'Extract all information needed to write a "{doc_type}" from these input documents.\n\n'
                f"{combined}\n\n"
                "Return JSON:\n"
                '{"system_name":"...","system_description":"...","technical_specs":{},'
                '"requirements":[],"infrastructure":{},"stakeholders":[],'
                '"key_facts":[],"other_details":{}}'
            ),
        },
    ], connection_id=connection_id)


def _generate_section(
    doc_type: str, section: dict, research: dict,
    style: str, reg_lang: list, connection_id: str = None,
) -> str:
    required = "\n".join(f"- {e}" for e in section.get("required_elements", []))
    reg = "\n".join(f"- {p}" for p in reg_lang[:10])
    return _llm_call([
        {
            "role": "system",
            "content": (
                "You are an expert technical writer for ISO 13485 regulated "
                "software V&V documentation."
            ),
        },
        {
            "role": "user",
            "content": (
                f'Write the "{section["heading"]}" section for a "{doc_type}" '
                f"document per ISO 13485.\n\n"
                f"Purpose: {section['description']}\n"
                f"Required elements:\n{required}\n\n"
                f"System info:\n{json.dumps(research, indent=2)[:5000]}\n\n"
                f"Style: {style}\n"
                f"Regulatory language:\n{reg}\n\n"
                "Write professional prose only — do not repeat the section heading."
            ),
        },
    ], connection_id=connection_id)


def _assemble_docx(doc_type: str, sections: list) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading(doc_type, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(
        f"Standard: ISO 13485  |  Generated: {datetime.now().strftime('%Y-%m-%d')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for heading, content in sections:
        doc.add_heading(heading, level=1)
        for block in content.split("\n\n"):
            if block.strip():
                doc.add_paragraph(block.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
